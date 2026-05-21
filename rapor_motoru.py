from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import io

def word_raporu_uret(df, s, secilen_dd):
    doc = Document()
    
    # 1. ANA BAŞLIK
    baslik = doc.add_heading('GEOTEKNİK DEĞERLENDİRME VE SIVILAŞMA ANALİZ RAPORU', 0)
    baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # --- BÖLÜM 1: STRATİGRAFİ ---
    doc.add_heading('1. Giriş ve Saha Stratigrafisi', level=1)
    
    zemin_tipleri = ", ".join(df['Zemin_Sinifi'].dropna().unique())
    doc.add_paragraph(
        f"Bu rapor, sahada gerçekleştirilen sondaj ve geoteknik etüt verilerine dayanılarak hazırlanmıştır. "
        f"İnceleme alanı zemin profili genel hatlarıyla {zemin_tipleri} tipi zemin tabakalarından oluşmaktadır. "
        f"Zemin katmanlarının dinamik davranışı ve yeraltı su seviyesinin (YASS) etkileri aşağıdaki bölümlerde detaylandırılmıştır."
    )

    # --- BÖLÜM 2: ZEMİN SINIFI VE SİSMİSİTE ---
    doc.add_heading('2. TBDY-2018 Yerel Zemin Sınıfı Değerlendirmesi', level=1)
    p_zemin = doc.add_paragraph()
    p_zemin.add_run("Türkiye Bina Deprem Yönetmeliği (TBDY-2018) Bölüm 16 kriterleri uyarınca yapılan değerlendirmede, saha verileri (")
    p_zemin.add_run(f"{s['zemin_nedeni']}").bold = True
    p_zemin.add_run(") dikkate alınmış olup, inceleme alanı yerel zemin sınıfı ")
    p_zemin.add_run(f"{s['zemin_sinifi']}").bold = True
    p_zemin.add_run(" olarak tayin edilmiştir. Sahanın sismik tasarım spektrumları bu sınıflandırma üzerinden türetilmiştir.")

    doc.add_paragraph(f"Tasarım Depremi (TBDY-2018): {secilen_dd}", style='List Bullet')
    doc.add_paragraph(f"PGA (En Büyük Yer İvmesi): {s['PGA']:.3f} g", style='List Bullet')
    doc.add_paragraph(f"Kısa Periyot Tasarım İvmesi (SDS): {s['SDS']:.3f} g", style='List Bullet')

    # --- BÖLÜM 3: SIVILAŞMA VE OTURMA YORUMU ---
    doc.add_heading('3. Sıvılaşma Potansiyeli ve Sismik Oturma Analizi', level=1)
    p_siv = doc.add_paragraph()
    p_siv.add_run(
        "Saha genelinde deprem etkisi altında kohezyonsuz zemin tabakalarında meydana gelebilecek sıvılaşma potansiyeli "
        "ve buna bağlı hacimsel şekil değiştirme analizleri gerçekleştirilmiştir. Yapılan analizler neticesinde saha genelinde hesaplanan maksimum "
        "sıvılaşma kaynaklı oturma miktarı "
    )
    p_siv.add_run(f"{s['max_oturma']:.2f} cm").bold = True
    p_siv.add_run(" olarak hesaplanmıştır. ")

    # Yapay Zeka Karar Mekanizması (Mühendislik Yorumu)
    if s['max_oturma'] > 4.0:
        run_risk = p_siv.add_run(
            "Hesaplanan bu değer, üst yapı güvenliği açısından yönetmeliklerde öngörülen tolere edilebilir sınırların (4.0 cm) üzerindedir. "
            "Sahada sıvılaşma riskine karşı zemin iyileştirmesi yapılması (jet-grout, taş kolon vb.) veya derin temel (kazık) "
            "sistemlerine geçilmesi statik ve geoteknik açıdan teknik bir zorunluluktur."
        )
        run_risk.font.color.rgb = RGBColor(200, 0, 0)
    else:
        run_risk = p_siv.add_run(
            "Hesaplanan bu değer, sığ temel sistemleri için öngörülen tolere edilebilir sınırlar içindedir. "
            "İncelenen profillerde üst yapı güvenliğini tehdit edecek düzeyde yıkıcı bir sıvılaşma riski öngörülmemektedir."
        )
        run_risk.font.color.rgb = RGBColor(0, 128, 0)

    # --- BÖLÜM 4: İYİLEŞTİRME (EĞER AKTİFSE) ---
    hedef_baslik_no = 4
    if s['toplam_kolon_sayisi'] > 0:
        doc.add_heading(f'{hedef_baslik_no}. Zemin İyileştirme Tasarımı (Baez Model)', level=1)
        doc.add_paragraph(
            f"Sıvılaşma kaynaklı kayma gerilmelerini (CSR) sönümlemek amacıyla sahada toplam {s['toplam_kolon_sayisi']} adet "
            f"rijit kolon planlanmıştır. %{s['Ar']*100:.1f} alan yer değiştirme oranına (Ar) sahip bu tasarım ile zemin taşıma gücü artırılmış "
            f"ve oturmalar tolere edilebilir seviyelere çekilmiştir."
        )
        hedef_baslik_no += 1

    # --- BÖLÜM 5: STATİK TASARIM TABLOSU ---
    doc.add_heading(f'{hedef_baslik_no}. Geoteknik Tasarım Parametreleri', level=1)
    doc.add_paragraph("Üstyapı statik projelendirmesinde (SAP2000, ideCAD vb.) kullanılmak üzere Kulhawy & Mayne (1990) ve Stroud (1974) ampirik bağıntılarıyla türetilen tabaka bazlı tasarım parametreleri aşağıda sunulmuştur:")
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    
    basliklar = ['Kuyu / Derinlik', 'Zemin', 'γ (kN/m³)', 'Cu (kPa)', 'ϕ (°)', 'E Modülü (kPa)']
    for i, baslik_metni in enumerate(basliklar):
        hdr_cells[i].text = baslik_metni
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True

    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = f"{row['Sondaj_No']} ({row['Derinlik_m']}m)"
        row_cells[1].text = str(row['Zemin_Sinifi'])
        row_cells[2].text = f"{row['Gamma_Tasarim']:.1f}"
        row_cells[3].text = f"{row['Cu_Tasarim']:.1f}" if pd.notna(row['Cu_Tasarim']) else "-"
        row_cells[4].text = f"{row['Phi_Acisi']:.1f}" if pd.notna(row['Phi_Acisi']) else "-"
        row_cells[5].text = f"{row['E_Modulu']:.0f}"

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()